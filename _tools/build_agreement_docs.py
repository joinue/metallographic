"""
Generate branded Microsoft Word (.docx) versions of the Distributor and
Reseller Agreements.

Re-uses the section data from build_agreements.py for a single source of truth.

Run from project root:  python _tools/build_agreement_docs.py
Output: downloads/pace-distributor-agreement.docx
        downloads/pace-reseller-agreement.docx
"""

from pathlib import Path
import re
import sys

sys.path.insert(0, str(Path(__file__).resolve().parent))

from build_agreements import DISTRIBUTOR_SECTIONS, RESELLER_SECTIONS

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


ROOT = Path(__file__).resolve().parent.parent
# Files live in an obfuscated subdirectory so a guessed URL like
# /downloads/pace-distributor-agreement.pdf returns 404. The partner-downloads
# page (which is password-gated) is the only place that knows the real path.
# To rotate the obfuscation: rename the directory below AND update the link
# paths in partner-downloads.html in lockstep.
DOWNLOADS = ROOT / "downloads" / "partner-b527ecc1"
DOWNLOADS.mkdir(parents=True, exist_ok=True)

PACE_BLUE = RGBColor(0x0A, 0x25, 0x40)
PACE_ACCENT = RGBColor(0x2D, 0x5A, 0x87)
MUTED = RGBColor(0x4F, 0x56, 0x6B)
LIGHT = RGBColor(0xA3, 0xAC, 0xBA)


def set_cell_border(cell, **kwargs):
    """Set borders on a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        if edge in kwargs:
            tag = OxmlElement(f"w:{edge}")
            for k, v in kwargs[edge].items():
                tag.set(qn(f"w:{k}"), v)
            tcBorders.append(tag)
    tcPr.append(tcBorders)


def add_header_footer(doc, title):
    """Add PACE-branded header and footer to every page."""
    section = doc.sections[0]

    # Header — top-left PACE name, top-right document title
    header = section.header
    header_para = header.paragraphs[0]
    header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = header_para.add_run("PACE Technologies Corporation")
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.bold = True
    run.font.color.rgb = PACE_BLUE
    tab_stops = header_para.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(6.5), WD_ALIGN_PARAGRAPH.RIGHT)
    run2 = header_para.add_run(f"\t{title}")
    run2.font.name = "Calibri"
    run2.font.size = Pt(9)
    run2.font.color.rgb = MUTED

    # Header rule
    header_rule = header.add_paragraph()
    p = header_rule._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "0A2540")
    pBdr.append(bottom)
    pPr.append(pBdr)

    # Footer — address (left) and page number (right)
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = footer_para.add_run("3601 E. 34th Street · Tucson, AZ 85713 · USA · +1 (520) 882-6598 · pace@metallographic.com")
    run.font.name = "Calibri"
    run.font.size = Pt(8)
    run.font.color.rgb = LIGHT
    tab_stops = footer_para.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(6.5), WD_ALIGN_PARAGRAPH.RIGHT)
    run_tab = footer_para.add_run("\t")
    # Page number field
    page_run = footer_para.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    page_run._r.append(fldChar1)
    page_run._r.append(instrText)
    page_run._r.append(fldChar2)
    page_run.font.name = "Calibri"
    page_run.font.size = Pt(8)
    page_run.font.color.rgb = LIGHT


def set_doc_defaults(doc):
    """Set body default font and margins."""
    section = doc.sections[0]
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.header_distance = Inches(0.4)
    section.footer_distance = Inches(0.4)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.25


def add_title_block(doc, title, subtitle, version, effective):
    """Top-of-document title block."""
    para = doc.add_paragraph()
    run = para.add_run(title)
    run.font.name = "Calibri"
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = PACE_BLUE
    para.paragraph_format.space_after = Pt(2)

    if subtitle:
        sub = doc.add_paragraph()
        srun = sub.add_run(subtitle)
        srun.font.name = "Calibri"
        srun.font.size = Pt(11)
        srun.font.color.rgb = MUTED
        sub.paragraph_format.space_after = Pt(2)

    meta = doc.add_paragraph()
    mrun = meta.add_run(f"Effective {effective}  ·  Template Version {version}")
    mrun.font.name = "Calibri"
    mrun.font.size = Pt(9)
    mrun.font.color.rgb = LIGHT
    meta.paragraph_format.space_after = Pt(18)

    # Divider rule
    rule_para = doc.add_paragraph()
    p = rule_para._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "E3E8EE")
    pBdr.append(bottom)
    pPr.append(pBdr)
    rule_para.paragraph_format.space_after = Pt(16)


def add_cover_block(doc, role: str):
    """Cover-style WHEREAS / parties block."""
    if role == "distributor":
        intro = (
            "This Distributor Agreement (\"Agreement\") is made and entered into "
            "as of [EFFECTIVE DATE], by and between PACE Technologies Corporation, "
            "an Arizona corporation having its principal place of business at "
            "3601 E. 34th Street, Tucson, Arizona 85713, USA (\"PACE\"), and "
            "[DISTRIBUTOR LEGAL NAME], a [JURISDICTION / ENTITY TYPE], having "
            "its principal place of business at [DISTRIBUTOR ADDRESS] "
            "(\"Distributor\")."
        )
        whereas = [
            "PACE is engaged in the manufacture and sale of metallographic equipment, consumables, and related products and services;",
            "Distributor desires to be appointed as a non-exclusive authorized distributor of such products in the territory identified below; and",
            "PACE is willing to appoint Distributor on the terms set forth in this Agreement.",
        ]
    else:
        intro = (
            "This Reseller Agreement (\"Agreement\") is made and entered into "
            "as of [EFFECTIVE DATE], by and between PACE Technologies Corporation, "
            "an Arizona corporation having its principal place of business at "
            "3601 E. 34th Street, Tucson, Arizona 85713, USA (\"PACE\"), and "
            "[RESELLER LEGAL NAME], a [JURISDICTION / ENTITY TYPE], having its "
            "principal place of business at [RESELLER ADDRESS] (\"Reseller\")."
        )
        whereas = [
            "PACE is engaged in the manufacture and sale of metallographic equipment, consumables, and related products and services;",
            "Reseller desires to be appointed as a non-exclusive authorized reseller of such products in the territory identified below; and",
            "PACE is willing to appoint Reseller on the terms set forth in this Agreement.",
        ]

    para = doc.add_paragraph()
    para.add_run(intro).font.size = Pt(11)
    para.paragraph_format.space_after = Pt(10)

    for w in whereas:
        p = doc.add_paragraph()
        bold = p.add_run("WHEREAS, ")
        bold.font.bold = True
        bold.font.color.rgb = PACE_BLUE
        p.add_run(w)
        p.paragraph_format.space_after = Pt(4)

    p = doc.add_paragraph()
    bold = p.add_run("NOW, THEREFORE, ")
    bold.font.bold = True
    bold.font.color.rgb = PACE_BLUE
    p.add_run(
        "in consideration of the mutual covenants and agreements set forth "
        "below and other good and valuable consideration, the receipt and "
        "sufficiency of which are acknowledged, the parties agree as follows:"
    )
    p.paragraph_format.space_after = Pt(16)

    # Territory fill-in
    t = doc.add_paragraph()
    tb = t.add_run("Territory: ")
    tb.font.bold = True
    tb.font.color.rgb = PACE_BLUE
    t.add_run("[NAMED TERRITORY — e.g., \"the country of [Country]\" or \"the European Economic Area excluding [excluded countries]\"]")
    t.paragraph_format.space_after = Pt(18)


def add_section_heading(doc, num, title):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(14)
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.keep_with_next = True

    num_run = para.add_run(f"{num}  ")
    num_run.font.name = "Calibri"
    num_run.font.size = Pt(10)
    num_run.font.bold = False
    num_run.font.color.rgb = LIGHT

    title_run = para.add_run(title.upper())
    title_run.font.name = "Calibri"
    title_run.font.size = Pt(12)
    title_run.font.bold = True
    title_run.font.color.rgb = PACE_BLUE


HTML_TAG_RE = re.compile(r"<[^>]+>")


def html_to_runs(para, html_fragment, base_size=Pt(11)):
    """Add runs to a paragraph from a small subset of HTML."""
    # Replace links: <a ...>X</a> -> X (keep visible text only in Word)
    text = re.sub(r'<a [^>]*>(.*?)</a>', r'\1', html_fragment, flags=re.DOTALL)
    # Handle <strong> bold blocks
    parts = re.split(r'(<strong>.*?</strong>)', text, flags=re.DOTALL)
    for part in parts:
        if part.startswith("<strong>"):
            inner = part[len("<strong>"):-len("</strong>")]
            inner = HTML_TAG_RE.sub("", inner)
            inner = unescape_entities(inner)
            r = para.add_run(inner)
            r.font.bold = True
            r.font.name = "Calibri"
            r.font.size = base_size
        else:
            clean = HTML_TAG_RE.sub("", part)
            clean = unescape_entities(clean)
            if clean:
                r = para.add_run(clean)
                r.font.name = "Calibri"
                r.font.size = base_size


def unescape_entities(s: str) -> str:
    return (
        s.replace("&amp;", "&")
        .replace("&lt;", "<")
        .replace("&gt;", ">")
        .replace("&quot;", '"')
        .replace("&#39;", "'")
        .replace("&nbsp;", " ")
    )


def render_body(doc, body_html):
    """Render the body HTML (paragraphs and lists) into the Word doc."""
    # Strip leading/trailing whitespace
    body = body_html.strip()

    # Match <p>...</p> and <ul>...</ul> blocks (non-greedy)
    block_re = re.compile(r"(<p>.*?</p>|<ul>.*?</ul>)", re.DOTALL)
    pos = 0
    for m in block_re.finditer(body):
        block = m.group(1)
        if block.startswith("<p>"):
            inner = block[len("<p>"):-len("</p>")].strip()
            para = doc.add_paragraph()
            para.paragraph_format.space_after = Pt(6)
            para.paragraph_format.line_spacing = 1.3
            html_to_runs(para, inner)
        elif block.startswith("<ul>"):
            inner = block[len("<ul>"):-len("</ul>")].strip()
            for li in re.findall(r"<li>(.*?)</li>", inner, flags=re.DOTALL):
                para = doc.add_paragraph(style="List Bullet")
                para.paragraph_format.space_after = Pt(3)
                para.paragraph_format.line_spacing = 1.25
                html_to_runs(para, li.strip())
        pos = m.end()


def render_summary(doc, summary_text):
    """Render the Summary callout as an italicized indented paragraph."""
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.2)
    para.paragraph_format.space_after = Pt(8)
    label = para.add_run("Summary  ")
    label.font.name = "Calibri"
    label.font.size = Pt(8)
    label.font.bold = True
    label.font.color.rgb = PACE_ACCENT
    label.font.all_caps = True

    body = para.add_run(summary_text)
    body.font.name = "Calibri"
    body.font.size = Pt(10)
    body.font.italic = True
    body.font.color.rgb = MUTED


def build_doc(role: str, sections, out_path: Path):
    doc = Document()
    set_doc_defaults(doc)

    title = "Distributor Agreement" if role == "distributor" else "Reseller Agreement"
    subtitle = "PACE Technologies Corporation — Template"

    add_header_footer(doc, title)
    add_title_block(doc, title, subtitle, version="2.0", effective="[EFFECTIVE DATE]")
    add_cover_block(doc, role)

    for sid, num, sec_title, summary, body in sections:
        add_section_heading(doc, num, sec_title)
        render_summary(doc, summary)
        render_body(doc, body)

    doc.save(str(out_path))
    print(f"  wrote {out_path.name} ({out_path.stat().st_size:,} bytes)")


def main():
    print("Building Distributor Agreement .docx…")
    build_doc(
        role="distributor",
        sections=DISTRIBUTOR_SECTIONS,
        out_path=DOWNLOADS / "pace-distributor-agreement.docx",
    )

    print("Building Reseller Agreement .docx…")
    build_doc(
        role="reseller",
        sections=RESELLER_SECTIONS,
        out_path=DOWNLOADS / "pace-reseller-agreement.docx",
    )

    print("Done.")


if __name__ == "__main__":
    main()
